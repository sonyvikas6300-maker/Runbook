import os
import json
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from groq import Groq

os.environ["GROQ_API_KEY"] = "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"
client = Groq()


# ---------------- NORMALIZE ---------------- #
def normalize(text):
    return text.lower().replace(" ", "").replace(".", "")


# ---------------- READ DOCX ---------------- #
def read_docx_sections(path):
    doc = Document(path)

    sections = {"bronze": [], "silver": [], "gold": []}
    current = None
    prev = ""

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        combined = normalize(prev + text)

        if "bronzelayer" in combined:
            current = "bronze"
            continue
        elif "silverlayer" in combined:
            current = "silver"
            continue
        elif "goldlayer" in combined:
            current = "gold"
            continue

        if current:
            sections[current].append(text)

        prev = text

    return (
        "\n".join(sections["bronze"]),
        "\n".join(sections["silver"]),
        "\n".join(sections["gold"])
    )


# ---------------- CLEAN SQL ---------------- #
def clean_sql(text):
    return "\n".join([
        l.strip() for l in text.split("\n")
        if l.strip() and "page" not in l.lower()
    ])


# ---------------- MINIMIZE METADATA ---------------- #
def minimize_metadata(meta):
    if not meta:
        return {}

    return {
        "database": meta.get("database"),
        "schema": meta.get("schema"),
        "table": meta.get("table"),
        "columns": meta.get("columns", [])[:50],
        "pipes": meta.get("pipes", [])[:5],
        "tasks": meta.get("tasks", [])[:5],
        "procedures": meta.get("procedures", [])[:5]
    }


# ---------------- AI (CHUNK) ---------------- #
def update_single_layer_sql(sql, metadata):
    prompt = f"""
You are a Snowflake SQL expert.

Update SQL using metadata.

Rules:
- Use correct database.schema.table
- Use only available columns
- Keep logic same

METADATA:
{json.dumps(metadata, indent=2)}

SQL:
{sql}

Return only SQL.
"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )

    return response.choices[0].message.content.strip()


# ---------------- DOCX HELPERS ---------------- #
def delete_paragraph(p):
    p._element.getparent().remove(p._element)


def insert_after(p, text):
    new_p = OxmlElement("w:p")
    p._element.addnext(new_p)
    para = Paragraph(new_p, p._parent)
    para.add_run(text)
    return para


# ---------------- REPLACE DOC ---------------- #
def replace_content(doc, b, s, g):
    paras = doc.paragraphs

    def find(keyword):
        for i, p in enumerate(paras):
            if keyword in p.text.lower().replace(" ", ""):
                return i
        return None

    bi = find("bronzelayer")
    si = find("silverlayer")
    gi = find("goldlayer")

    def replace(start, end, content):
        for i in range(start + 1, end):
            delete_paragraph(paras[i])
        insert_after(paras[start], content)

    replace(bi, si, b)
    replace(si, gi, s)
    replace(gi, len(paras), g)

    return doc
